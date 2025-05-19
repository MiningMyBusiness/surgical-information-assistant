import streamlit as st
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from utils.agents import orchestrator, DeRetSynState, handle_simple_question
import torch
import sys
from docx import Document
import io

## Run with: streamlit run app.py local async -- --server.port 8501 --server.address 0.0.0.0

# Check if the "local" argument is provided
# local = len(sys.argv) > 1 and sys.argv[1] == "local"
key_prefix = ""
# if local:
#     key_prefix = "local_"

# Check if the "async" argument is provided
# RUN_ASYNC = len(sys.argv) > 2 and sys.argv[2].lower() == "async"
RUN_ASYNC = True

# fix issue with torch path, in case it occurs with streamlit
torch.classes.__path__ = []

# Set up LLM
llm = ChatOpenAI(model=st.secrets[key_prefix + "model_2_small"],
                 api_key=st.secrets[key_prefix + "api_key_2"],
                 base_url=st.secrets[key_prefix + "base_url_2"])

# --- Decision Layer --- #
decision_prompt = PromptTemplate.from_template(
    """Determine if the following question requires access to specific medical documents related to detailed surgical information to be answered accurately.
Think step-by-step and reason through your answer. Respond in the following format: 

<thinking> Your reasoning here... </thinking>
<answer> yes OR no </answer>

Here are some examples:

Question:
What is the primary purpose of the coronary artery bypass graft?
Response:
<thinking> The question is asking about a coronary bypass graft which is related to surgery. So yes, this question is about surgery. </thinking>
<answer> yes </answer>

Question:
Is machine learning useful for solving complex medical problems?
Response:
<thinking> The question is asking about machine learning as a tool for solving complex medical problems. While the question does specify medical problems, it does not refer to surgery or topics related to surgery. </thinking>
<answer> no </answer>

Question:
Can you elaborate on your previous response about suturing or rephrase it?
Response:
<thinking> The question is asking for an explanation of suturing that was given earlier in the conversation. While the question does mention suturing which is related to surgery, it is not asking about suturing specifically but rather requesting to explain a previous response. </thinking>
<answer> no </answer>

Question:
What are some core differences between robotic and laparasopic inguinal hernia?
Response:
<thinking> The question is asking about the differences between robotic and laparosopic inguinal hernia. Both robotic and laparosopic surgery involve surgery on the inguinal hernia. This is clearly a question about surgery and requires investigation into documentation to answer. </thinking>
<answer> yes </answer>

Question:
{question}
Response:
"""
)

def is_about_surgery(question):
    response = llm.invoke(decision_prompt.format(question=question)).content.strip()
    decision = response.split("<answer>")[1].split("</answer>")[0].strip().lower()
    return "yes" in decision

def update_user_input_with_context(user_input, chat_history):
    # Get the last 4 turns of the chat history
    recent_history = chat_history[-3:]
    
    # Format the chat history for the LLM
    formatted_history = "\n".join([f"User: {turn['user']}\nAssistant: {turn['bot']}" for turn in recent_history])
    
    # Prepare the prompt for the LLM
    prompt = f"""Given the following chat history and a new user input, update the user input to include any relevant context from the chat history. Only add necessary context, and keep the updated input concise.
Think step-by-step and reason through your answer. Respond in the following format:

<thinking> Your reasoning here... </thinking>
<updated_input> Updated user input here... </updated_input>
    
Here are some examples:

CHAT HISTORY:
User: What are the main types of sutures used in surgery?
Assistant: The main types of sutures used in surgery are absorbable and non-absorbable sutures. Absorbable sutures dissolve over time, while non-absorbable sutures need to be removed. Common absorbable sutures include catgut and synthetic materials like polyglycolic acid. Non-absorbable sutures are often made of materials like nylon or silk.
User: How long do absorbable sutures typically last?
Assistant: The duration of absorbable sutures varies depending on the material used. Generally, they can last anywhere from a few days to several months. For example, fast-absorbing sutures may dissolve within 7-10 days, while slower-absorbing sutures can take 60-90 days or even longer.
NEW USER INPUT:
What about non-absorbable ones?
RESPONSE:
<thinking> Considering the previous discussion the user is referring to non-absorbable sutures. Specifically, they are asking about the lifespan or removal time-frame for non-absorbable sutures. I'll restate the question to include this context from the chat history. </thinking>
<updated_input> What is the typical lifespan or removal timeframe for non-absorbable sutures used in surgery? </updated_input>



CHAT HISTORY:
User: What is a laparotomy?
Assistant: A laparotomy is a surgical procedure that involves making a large incision in the abdomen to gain access to the abdominal cavity. It allows surgeons to examine the abdominal organs, diagnose problems, and perform necessary treatments. Laparotomies are used for various purposes, including exploratory surgery, treatment of abdominal trauma, and removal of tumors.
User: Are there any alternatives to this procedure?
Assistant: Yes, there are alternatives to laparotomy, particularly minimally invasive techniques. The main alternative is laparoscopy, also known as keyhole surgery. In laparoscopy, several small incisions are made instead of one large incision. A camera and specialized instruments are inserted through these small incisions to perform the surgery. This technique often results in less pain, faster recovery, and smaller scars compared to traditional laparotomy.

NEW USER INPUT:
What are the risks?

RESPONSE:
<thinking> The user is asking about risks, but it's not clear whether they're referring to laparotomy or laparoscopy, both of which were discussed in the previous messages. Since laparotomy was the original topic and laparoscopy was introduced as an alternative, it would be helpful to ask about the risks of both procedures for a comprehensive answer. </thinking>
<updated_input> What are the risks associated with both laparotomy and laparoscopy procedures? </updated_input>



CHAT HISTORY:
User: What are the potential complications of laparoscopic surgery?
Assistant: Laparoscopic surgery can have several potential complications, including:
1. Infection: Laparoscopic surgery can lead to infections, especially in patients with infections that can be spread through the abdominal cavity.
2. Stool bleeding: Laparoscopic surgery can lead to stool bleeding, especially if the surgical instrument is inserted directly into the abdominal cavity.
3. Pain: Laparoscopic surgery can be painful, especially for patients with pre-existing conditions or those who have had previous laparoscopic surgeries.

NEW USER INPUT:
Can you think of any more?

RESPONSE:
<thinking> The user is asking if there are any more complications, but it's important to note that the original question was about laparoscopic surgery, not about potential complications. I'll restate the question to clarify that the original topic was laparoscopic surgery. </thinking>
<updated_input> Are there any more potential complications associated with laparoscopic surgery besides infection, stool bleeding, and pain? </updated_input>



CHAT HISTORY:
{formatted_history}

New User Input:
{user_input}

RESPONSE:
"""

    # Make the LLM call
    response = llm.invoke(prompt).content.strip()
    answer = response.split("<updated_input>")[1].split("</updated_input>")[0].strip()
    return answer

# --- LangGraph Invocation --- #
def run_agents(user_input):
    initial_state = DeRetSynState(
        original_question=user_input,
        answers="",
        iterations=0,
        faiss_index_path="surgical_faiss_index",
        model=st.secrets[key_prefix + "model_2_small"],
        api_key=st.secrets[key_prefix + "api_key_2"],
        base_url=st.secrets[key_prefix + "base_url_2"],
        verbose=True,
        run_async=RUN_ASYNC
    )

    progress_placeholder = st.empty()
    result_placeholder = st.empty()
    progress_placeholder_2 = st.empty()

    with st.spinner("Processing your question..."):
        progress_placeholder.text("Decomposing your question into sub-questions...")
        for result in orchestrator(initial_state):
            if result["step"] == "decompose_complete":
                progress_placeholder.text("Decomposed your question into the following sub-questions.")
                result_placeholder.text("Generated sub-questions:\n- " + "\n- ".join(result["sub_questions"]))
                progress_placeholder_2.text("Retrieving relevant information for each sub-question...")
            
            elif result["step"] == "retrieve_complete":
                progress_placeholder.text("Retrieved relevant information.")
                result_placeholder.text("")
                progress_placeholder_2.text("Synthesizing information...")
            
            elif result["step"] == "synthesize_complete":
                if not result["done"]:
                    progress_placeholder.text("Synthesized information.")
                    result_placeholder.text("Determined that more information is needed. Generated new queries:\n-" + "\n-".join(result["new_queries"]))
                    progress_placeholder_2.text("Retrieving additional information...")
                else:
                    progress_placeholder.text("Synthesized information.")
                    result_placeholder.text("")
                    progress_placeholder_2.text("Processing complete.")
            
            elif result["step"] == "start_best_effort":
                progress_placeholder_2.text("Facing difficulty completely answering the question with available documents. Generating best effort answer using help from Wikipedia...")
            
            elif result["step"] == "best_effort_complete":
                progress_placeholder.text("Best effort answer generated.")
                result_placeholder.text(f"Wikipedia context used for best effort answer:\n{result['wiki_results']}")
            
            elif result["step"] == "final":
                progress_placeholder.text("Final answer generated.")
                result_placeholder.text("")
                progress_placeholder_2.text(f"Processing complete. Number of iterations needed: {result['state']['iterations']}")

    return result["state"]


if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

def clear_chat_history():
    st.session_state.chat_history = []


def download_chat_history():
    doc = Document()
    doc.add_heading('Chat History', 0)
    
    for msg in st.session_state.chat_history:
        doc.add_paragraph(f"User: {msg['user']}")
        doc.add_paragraph(f"Assistant: {msg['bot']}")
        doc.add_paragraph("\n")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# --- Streamlit UI --- #
st.set_page_config(layout="wide")

st.title("📚 Surgical Information Assistant")
st.info("""This app uses Agents and RAG to provide intelligent 
responses to your questions about surgery by using the [Open Manual of Surgery in Resource-Limited Settings](https://www.vumc.org/global-surgical-atlas/about)
created by the Vanderbilt University Medical Center.
DISCLAIMER: This app, the code, and the LLM
responses are not endorsed by or associated with the Vanderbilt University Medical Center. 
The code is open-source 😊 
and follows the same licensing as the documents with the Open Manual of Surgery (CC-v1.0 Universal).
Answers are generated by an LLM and LLMs can make mistakes. Double-check it. Backend is powered by Llama3.1-8b and Llama3.2-3b models.
""")

col1, col2 = st.columns(2)
with col1:
    if st.button("Clear chat"):
        clear_chat_history()
        st.rerun()
with col2:
    with st.container():
        st.markdown(
            """
            <style>
            .stDownloadButton {
                float: right;
            }
            </style>
            """,
            unsafe_allow_html=True
        )
        if st.download_button(
            label="Download Chat History",
            data=download_chat_history(),
            file_name="chat_history.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            st.success("Chat history downloaded successfully!")

# if RUN_ASYNC:
#     st.info("Running in async mode")
# else:
#     st.info("Running in sync mode")

# Show past chat messages
for i,msg in enumerate(st.session_state.chat_history):
    st.chat_message("user").write(msg["user"])
    with st.chat_message("assistant"):
        st.markdown(msg["bot"])
        answers = msg.get("answers", "")
        wiki_results = msg.get("wiki_results", "")
        pending_queries = msg.get("pending_queries", [])
        with st.expander("View reference information used to generate the answer"):
            if wiki_results:
                st.write("Wikipedia context used to answer the question:\n" + wiki_results)
            else:
                st.write("Document QA used to answer the question:\n" + answers)
        if pending_queries and i == len(st.session_state.chat_history) - 1:
            with st.expander("Suggested follow-up questions"):
                st.write("\n- " + "\n- ".join(pending_queries))


default_text = "Ask me something about about surgery."
if user_input := st.chat_input(default_text):
    with st.chat_message("user"):
        st.write(user_input)

    if st.session_state.chat_history:
        updated_input = update_user_input_with_context(user_input, st.session_state.chat_history)
    else:
        updated_input = user_input

    if is_about_surgery(updated_input):
        response = run_agents(updated_input)
    else:
        response = handle_simple_question(user_input, st.session_state.chat_history, llm)
    with st.chat_message("assistant"):
        st.markdown(response["final_answer"])
        answers = response.get("answers", "")
        wiki_results = response.get("wiki_results", "")
        pending_queries = response.get("pending_queries", [])
        with st.expander("View reference information used to generate the answer"):
            if wiki_results:
                st.write("Wikipedia context used to answer the question:\n" + wiki_results)
            else:
                st.write("Document QA used to answer the question:\n" + answers)
        if pending_queries:
            with st.expander("Suggested follow-up questions"):
                st.write("\n- " + "\n- ".join(pending_queries))
        st.session_state.chat_history.append({"user": user_input, "bot": response["final_answer"],
                                                "answers": response.get("answers", ""),
                                                "wiki_results": response.get("wikipedia_results", None)})
